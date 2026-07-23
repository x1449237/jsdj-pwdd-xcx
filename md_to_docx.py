import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if level == 1:
        run.font.size = Pt(22)
    elif level == 2:
        run.font.size = Pt(18)
    elif level == 3:
        run.font.size = Pt(15)
    else:
        run.font.size = Pt(13)
    return p


def add_paragraph(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 * (level + 1))
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(11)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_bg(cell, '2E86AB')

    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.size = Pt(10)
            if r_idx % 2 == 1:
                set_cell_bg(cell, 'F0F7FA')

    doc.add_paragraph()
    return table


def parse_inline_markdown(text):
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    return text


def md_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5

    i = 0
    in_code_block = False
    code_content = []

    while i < len(lines):
        line = lines[i].rstrip('\n')

        if line.startswith('```'):
            if in_code_block:
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_content))
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_content.append(line)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if line.startswith('# '):
            add_heading(doc, line[2:].strip(), 1)
        elif line.startswith('## '):
            add_heading(doc, line[3:].strip(), 2)
        elif line.startswith('### '):
            add_heading(doc, line[4:].strip(), 3)
        elif line.startswith('#### '):
            add_heading(doc, line[5:].strip(), 4)
        elif line.startswith('---'):
            p = doc.add_paragraph()
            p.add_run('_' * 50)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1

            if len(table_lines) >= 2:
                headers = [c.strip() for c in table_lines[0].strip('|').split('|')]
                rows = []
                for tl in table_lines[2:]:
                    if '---' in tl:
                        continue
                    cells = [c.strip() for c in tl.strip('|').split('|')]
                    cells = [parse_inline_markdown(c) for c in cells]
                    rows.append(cells)
                add_table(doc, headers, rows)
            continue
        elif line.startswith('- **') or line.startswith('- '):
            level = 0
            stripped = line
            while stripped.startswith('  '):
                level += 1
                stripped = stripped[2:]
            text = stripped[2:] if stripped.startswith('- ') else stripped[3:]
            text = parse_inline_markdown(text)
            add_bullet(doc, text, level)
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            text = parse_inline_markdown(text)
            p = doc.add_paragraph(style='List Number')
            run = p.add_run(text)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.size = Pt(11)
        else:
            text = parse_inline_markdown(line)
            add_paragraph(doc, text)

        i += 1

    doc.save(docx_path)
    print(f'已生成: {docx_path}')


if __name__ == '__main__':
    md_to_docx(
        '/workspace/平台优化升级全方案_V2.0.md',
        '/workspace/戟三电竞平台优化升级全方案_V2.0.docx'
    )
